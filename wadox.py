#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Feb 28 16:46:44 2025

@author: Fslucidi
"""


from vosk import Model, KaldiRecognizer
import wave
from docx import Document

# Path to the model and audio file
model_path = "/path/to/vosk-model"
audio_file_path = "/path/to/input_audio.wav"
output_word_file = "/path/to/output_transcription.docx"

def transcribe_and_save(audio_path, output_docx):
    model = Model(model_path)

    with wave.open(audio_path, "rb") as wf:
        recognizer = KaldiRecognizer(model, wf.getframerate())

        print("Transcription in progress...")
        transcript = []
        while True:
            data = wf.readframes(4000)
            if len(data) == 0:
                break
            if recognizer.AcceptWaveform(data):
                result = recognizer.Result()
                transcript.append(result)

        final_result = recognizer.FinalResult()
        transcript.append(final_result)
        print("Transcription completed.")

    # Save the transcription to a Word file
    doc = Document()
    doc.add_heading("Audio Transcription", level=1)
    
    for segment in transcript:
        doc.add_paragraph(segment)

    doc.save(output_docx)
    print(f"Transcription saved in {output_docx}")

# Execute transcription and save to Word
transcribe_and_save(audio_file_path, output_word_file)

# ONLINE METHOD: BUT HAS FILE SIZE LIMIT

# import os
# from pydub import AudioSegment

# # Add FFmpeg to Python's PATH
# os.environ["PATH"] += os.pathsep + "/path/to/ffmpeg"

# # Verify FFmpeg is now recognized
# from pydub.utils import which
# print("ffmpeg path:", which("ffmpeg"))
# print("ffprobe path:", which("ffprobe"))

# import speech_recognition as sr

# def convert_to_wav(audio_path):
#     # Convert the audio file to WAV format if not already
#     audio = AudioSegment.from_file(audio_path)
#     wav_path = audio_path.rsplit('.', 1)[0] + ".wav"
#     audio.export(wav_path, format="wav")
#     return wav_path

# def transcribe_audio(audio_path):
#     recognizer = sr.Recognizer()

#     # Convert to WAV if necessary
#     if not audio_path.endswith('.wav'):
#         audio_path = convert_to_wav(audio_path)

#     with sr.AudioFile(audio_path) as source:
#         print("Transcription in progress...")
#         audio_data = recognizer.record(source)
#         try:
#             transcription = recognizer.recognize_google(audio_data, language="it-IT")
#             print("Transcription completed:")
#             print(transcription)
#         except sr.UnknownValueError:
#             print("Speech was not recognized.")
#         except sr.RequestError:
#             print("Error in the speech recognition service request.")